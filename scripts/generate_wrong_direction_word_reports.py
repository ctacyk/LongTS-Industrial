# -*- coding: utf-8 -*-
"""生成方向错误类型分析报告（最强模型 + benchmark 全模型汇总）Word 文档。"""
from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

SCRIPTS_DIR = Path(__file__).resolve().parent
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from analyze_wrong_direction_types import compute_stats_from_wrongs, load_wrong_samples

REPO = Path(__file__).resolve().parents[1]
BENCH = REPO / "results" / "benchmark_eval_v2"
OUT_DIR = REPO / "output" / "doc"


def _set_doc_defaults(doc: Document) -> None:
    sec = doc.sections[0]
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val


def _type_order(stats: dict) -> list[str]:
    pc = stats["primary_type_counts"]
    return sorted(pc.keys(), key=lambda k: (-pc[k], k))


def build_doc(
    title: str,
    subtitle_lines: list[str],
    stats: dict,
    eval_results_path: Path | None,
    extra_notes: list[str] | None = None,
) -> Document:
    doc = Document()
    _set_doc_defaults(doc)
    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in subtitle_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(10)
    doc.add_paragraph()

    doc.add_heading("一、统计口径与分类方法", level=1)
    doc.add_paragraph(
        "本报告基于 LongTS-Industrial Benchmark 评测输出中的 "
        "`eval_detailed.jsonl`。方向错误定义为 LLM Judge 给出的 "
        "`direction == \"wrong\"`（核心结论与标准答案矛盾，方向对齐得分 DA=0）。"
    )
    doc.add_paragraph(
        "错误类型由对 Judge 中文 `reason` 字段的规则分类得到：规则按优先级依次匹配，"
        "每个样本仅归入一个主类；另给出多标签线索占比（词条可重叠，仅供辅助解读）。"
        "对主类为「多种原因混合」的样本，使用 `classify_other_detail` 再做互斥子类划分。"
        "该分类属工程化归纳，非数据集人工分层标注。"
    )

    if eval_results_path and eval_results_path.is_file():
        with open(eval_results_path, encoding="utf-8") as f:
            er = json.load(f)
        doc.add_heading("二、模型总体评测指标（来自 eval_results.json）", level=1)
        ov = er["overall"]
        dd = er.get("direction_dist", {})
        lines = [
            f"模型名称：{er.get('model', '')}",
            f"总样本数：{er.get('total_samples', ov['n'])}",
            f"HES：{ov['hes']:.4f}（综合得分）",
            f"DA：{ov['da']:.4f}（方向对齐均值）",
            f"SS：{ov['ss']:.4f}（语义相似度）",
            f"AQ：{ov['aq']:.4f}（答案质量）",
            f"方向分布 - aligned：{dd.get('aligned', '-')}, partial：{dd.get('partial', '-')}, "
            f"wrong：{dd.get('wrong', '-')}",
            f"本报告方向错误样本数 N：{stats['wrong_n']}（与 direction_dist.wrong 一致）",
        ]
        for line in lines:
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_heading("二、样本概况", level=1)
        doc.add_paragraph(f"方向错误样本数 N：{stats['wrong_n']}")

    n = max(stats["wrong_n"], 1)
    order = _type_order(stats)

    doc.add_heading("三、主错误类型分布（互斥）", level=1)
    rows = []
    for k in order:
        v = stats["primary_type_counts"][k]
        rows.append([k, str(v), f"{100.0 * v / n:.2f}%"])
    _add_table(doc, ["错误类型（主类）", "数量", "占方向错误比例"], rows)

    osub = stats.get("other_subtype_counts") or {}
    if osub:
        doc.add_heading("3.1 「多种原因混合」二次细分（仅统计该主类内部）", level=2)
        doc.add_paragraph(
            "下列比例分母为「多种原因混合」样本数，不是全部方向错误样本数。"
        )
        on = max(int(stats.get("other_n", 0)), 1)
        sub_rows = []
        for k, v in sorted(osub.items(), key=lambda x: -x[1]):
            sub_rows.append([k, str(v), f"{100.0 * v / on:.2f}%"])
        _add_table(doc, ["细分标签", "数量", "占多种原因混合比例"], sub_rows)

        om = stats.get("other_mix_opening_counts") or {}
        if om:
            doc.add_heading("3.2 「多种原因混合」中混合句式句首粗分桶（辅助）", level=2)
            doc.add_paragraph(
                "句首统计仅用于快速观察 Judge 评语习惯，同一桶内仍可能包含多种真实错误形态。"
            )
            mn = max(int((stats.get("other_subtype_counts") or {}).get("其他子类-混合或低频句式", 0)), 1)
            mor = []
            for k, v in sorted(om.items(), key=lambda x: -x[1]):
                mor.append([k, str(v), f"{100.0 * v / mn:.2f}%"])
            _add_table(doc, ["混合句首类型", "数量", "占混合句式子类比例"], mor)

    doc.add_heading("四、按领域（domain）", level=1)
    headers = ["领域", "方向错误数"] + order
    drows = []
    for dom in sorted(stats["by_domain"].keys()):
        dc = stats["by_domain"][dom]
        total = sum(dc.values())
        drows.append([dom, str(total)] + [str(dc.get(t, 0)) for t in order])
    _add_table(doc, headers, drows)

    doc.add_heading("五、按难度层级（level）", level=1)
    headers = ["Level", "方向错误数"] + order
    lrows = []
    for lv in sorted(stats["by_level"].keys()):
        lc = stats["by_level"][lv]
        total = sum(lc.values())
        lrows.append([lv, str(total)] + [str(lc.get(t, 0)) for t in order])
    _add_table(doc, headers, lrows)

    doc.add_heading("六、按题目难度（difficulty）", level=1)
    headers = ["difficulty", "方向错误数"] + order
    erows = []
    for df in sorted(stats["by_difficulty"].keys()):
        fc = stats["by_difficulty"][df]
        total = sum(fc.values())
        erows.append([df, str(total)] + [str(fc.get(t, 0)) for t in order])
    _add_table(doc, headers, erows)

    doc.add_heading("七、多标签线索占比（可重叠）", level=1)
    mt = stats["multi_tag_counts"]
    mrows = [[k, str(v), f"{100.0 * v / n:.2f}%"] for k, v in sorted(mt.items(), key=lambda x: -x[1])]
    _add_table(doc, ["线索标签", "出现次数", "占方向错误比例"], mrows)

    if extra_notes:
        doc.add_heading("附录说明", level=1)
        for note in extra_notes:
            doc.add_paragraph(note)

    doc.add_paragraph()
    foot = doc.add_paragraph(f"生成日期：{date.today().isoformat()}")
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc


def discover_eval_detailed() -> list[Path]:
    paths = sorted(BENCH.rglob("eval_detailed.jsonl"))
    out = []
    for p in paths:
        sp = str(p)
        if "backup" in sp.lower():
            continue
        out.append(p)
    return out


def _safe_docx_filename(name: str) -> str:
    bad = '<>:"/\\|?*'
    for c in bad:
        name = name.replace(c, "_")
    return name.strip() or "unknown_model"


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    eval_paths = discover_eval_detailed()
    if not eval_paths:
        raise FileNotFoundError("未发现任何 eval_detailed.jsonl")

    written: list[Path] = []
    for eval_path in eval_paths:
        model_dir = eval_path.parent
        name = model_dir.name
        wrongs = load_wrong_samples(str(eval_path))
        stats = compute_stats_from_wrongs(wrongs)
        er_path = model_dir / "eval_results.json"
        doc = build_doc(
            title="Benchmark 方向错误类型分析报告",
            subtitle_lines=[f"评测目录：{name}"],
            stats=stats,
            eval_results_path=er_path if er_path.is_file() else None,
            extra_notes=None,
        )
        fname = f"方向错误类型分析_{_safe_docx_filename(name)}.docx"
        out_path = OUT_DIR / fname
        doc.save(str(out_path))
        written.append(out_path)
        print("Written:", out_path)

    print(f"共生成 {len(written)} 个 Word 报告（每个模型一份）。")


if __name__ == "__main__":
    import os

    os.chdir(REPO)
    main()
